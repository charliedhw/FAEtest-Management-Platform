# -*- coding: utf-8 -*-
"""生成测试项目管理平台 开发设计部署方案文档"""
import sys
sys.stdout.reconfigure(encoding="utf-8", errors="replace")
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# 全局中文字体
def set_cn_font(run, name="宋体", size=10.5, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), name)

def add_heading(text, level=1):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    sizes = {1: 18, 2: 15, 3: 13, 4: 12}
    run.font.name = "微软雅黑"
    run.font.size = Pt(sizes.get(level, 12))
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x1E, 0x3A) if level == 1 else RGBColor(0x1A, 0x2A, 0x6C)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_para(text, size=10.5, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_cn_font(run, size=size, bold=bold)
    return p

def add_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        run = hdr[i].paragraphs[0].add_run(h)
        set_cn_font(run, name="微软雅黑", size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1A2A6C')
        hdr[i]._tc.get_or_add_tcPr().append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            run = cells[i].paragraphs[0].add_run(str(val))
            set_cn_font(run, size=9.5)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    # 灰色底纹
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    return p

# ===================== 封面 =====================
for _ in range(4):
    doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("测试项目管理平台")
set_cn_font(run, name="微软雅黑", size=32, bold=True, color=RGBColor(0xC0, 0x1E, 0x3A))

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("开发设计部署方案")
set_cn_font(run, name="微软雅黑", size=22, bold=True, color=RGBColor(0x1A, 0x2A, 0x6C))

doc.add_paragraph()
tag = doc.add_paragraph()
tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tag.add_run("—— 替换邮件申请与在线表格的一体化测试管理平台 ——")
set_cn_font(run, size=12, color=RGBColor(0x88, 0x88, 0x88))

for _ in range(8):
    doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run("中科曙光  战略客户服务部\n\n2026 年 8 月")
set_cn_font(run, name="微软雅黑", size=14, color=RGBColor(0x33, 0x33, 0x33))

doc.add_page_break()

# ===================== 1 项目概述 =====================
add_heading("一、项目概述", 1)

add_heading("1.1 建设背景", 2)
add_para("公司测试业务原先采用邮件方式进行测试申请与审批，测试日报、周报也通过邮件发送，数据统计依赖飞书在线表格手工维护。该方式存在诸多痛点：", indent=True)
for t in [
    "申请无统一入口，进度不可查，易漏件、无留痕；",
    "审批意见分散在邮件中，驳回原因无记录，超期审批无系统控制；",
    "日报周报靠人工汇总，无法按项目沉淀测试进展；",
    "数据统计口径不统一，状态更新滞后，无法自动产出报表；",
    "测试资源与借用靠人工跟踪，超期风险高、费用归属不清。",
]:
    add_para("• " + t, indent=True)
add_para("为解决上述问题，建设测试项目管理平台，覆盖“申请→审批→资源→执行→报告→统计”全生命周期，替换原有邮件申请与在线表格模式，并预留与 OA 用户认证对接能力。", indent=True)

add_heading("1.2 建设目标", 2)
for t in [
    "统一测试申请入口，售前在线提交，全流程线上流转、节点留痕；",
    "多级审批（测试审批组→分管领导）线上化，支持驳回、超时提醒；",
    "测试任务由 FAE 负责人统一分配，测试人员在线填写进展、上传报告；",
    "资源池与借用闭环管理，超期自动预警；",
    "多维度统计分析看板，一键导出与原有口径一致的 Excel 报表；",
    "按角色的数据权限隔离，保障数据安全。",
]:
    add_para("• " + t, indent=True)

add_heading("1.3 平台访问信息", 2)
add_table(
    ["项目", "内容"],
    [
        ["访问地址", "http://192.168.101.165"],
        ["部署服务器", "192.168.101.165（Rocky Linux 8.10 虚拟机）"],
        ["部署方式", "Docker + Docker Compose 容器化部署"],
        ["管理员账号", "admin / Admin@123"],
        ["迁移历史数据", "175 个测试项目、569 条测试进展记录"],
    ],
    widths=[4, 12]
)

# ===================== 2 总体架构 =====================
add_heading("二、总体逻辑架构", 1)
add_para("平台采用前后端分离的分层架构，模块化单体设计，预留认证与外部系统对接能力。", indent=True)
add_code(
"""┌──────────────────────────────────────────────────────────┐
│  展现层   Web端(PC)   Vue3 + Element Plus + ECharts         │
│  项目统计│测试申请│审批中心│项目清单│资源管理│系统管理         │
├──────────────────────────────────────────────────────────┤
│  接入层   Nginx 反向代理 / 统一鉴权(JWT)                     │
├──────────────────────────────────────────────────────────┤
│  业务服务层  (Spring Boot 3 模块化单体)                      │
│   流程服务    申请单状态机、多级审批、驳回/撤回、超时提醒       │
│   项目服务    项目清单、进展日报、重点项目、按钮级权限          │
│   资源服务    资源池、占用/归还、借用规则、超期预警             │
│   报告服务    测试报告上传/版本/下载(MinIO)                   │
│   统计服务    看板、多维统计、自定义维度、Excel导出            │
│   消息服务    站内信(可点击跳转) + 邮件(预留)                 │
│   系统服务    用户/角色/用户组/字典/权限/审计/定时任务         │
├──────────────────────────────────────────────────────────┤
│  数据层   MySQL(业务库) │ Redis(会话/缓存) │ MinIO(对象存储)  │
├──────────────────────────────────────────────────────────┤
│  集成层(预留)  OA认证(OAuth2/CAS) │ SPM系统 │ 邮件/企微通知    │
└──────────────────────────────────────────────────────────┘""")

add_heading("2.1 关键架构决策", 2)
add_table(
    ["决策点", "选型", "说明"],
    [
        ["架构模式", "模块化单体", "用户规模百人级、流程集中，单体降低运维成本，模块边界清晰便于后续拆分"],
        ["流程引擎", "轻量状态机自研", "流程固定且仅两级审批，状态+节点+审批人配置即可，无需引入重型工作流"],
        ["认证体系", "本地账号+JWT，AuthProvider 抽象", "预留 oa_account 字段与认证适配器，二期平滑切换 OA 认证"],
        ["借用规则", "规则参数化", "30/90/365 天阈值、2% 比例等全部入库可配，制度调整不改代码"],
        ["文件存储", "MinIO 私有化对象存储", "测试报告大文件，免 NAS 依赖"],
    ],
    widths=[3, 5, 8]
)

# ===================== 3 用户与权限 =====================
add_heading("三、用户、角色与权限设计", 1)

add_heading("3.1 角色体系（RBAC）", 2)
add_table(
    ["角色编码", "角色名称", "职责"],
    [
        ["ADMIN", "系统管理员", "全部权限，系统维护"],
        ["SALES", "销售", "查看关联自己的测试项目，更新中标信息"],
        ["PRESALES", "售前工程师", "发起测试申请，跟踪自己提交/关联的项目"],
        ["APPROVER", "测试审批组", "审批测试申请，可驳回，可看全部数据"],
        ["LEADER", "分管领导", "超 90 天借测二级审批，可看全部数据"],
        ["TESTER", "FAE 测试工程师", "执行测试、填写进展、上传报告（仅限被分配项目）"],
        ["FAE_LEADER", "FAE 负责人", "分配测试任务给 FAE，可看全部测试数据"],
        ["RESOURCE_ADMIN", "资源管理员", "资源池与借用管理，可看全部数据"],
        ["BUSINESS", "商务", "借用跟催与超期处理"],
    ],
    widths=[3.2, 3.5, 9.3]
)

add_heading("3.2 用户组", 2)
add_para("平台引入用户组概念用于流程流转定位处理人，内置四个组，支持组成员与负责人维护：", indent=True)
add_table(
    ["组编码", "组名称", "负责人", "用途"],
    [
        ["SALES_GROUP", "销售组", "—", "销售归属"],
        ["PRESALES_GROUP", "售前组", "—", "售前归属"],
        ["APPROVER_GROUP", "测试审批组", "—", "审批流转"],
        ["FAE_GROUP", "FAE 测试工程师组", "戴海伟(daihw)", "任务分配目标人"],
    ],
    widths=[3.5, 4, 3.5, 5]
)

add_heading("3.3 数据权限", 2)
add_para("平台在申请单、项目清单、统计看板等各模块统一实施按角色的数据隔离：", indent=True)
add_table(
    ["角色", "可见数据范围"],
    [
        ["管理员 / 测试审批组 / 分管领导 / FAE 负责人 / 资源管理员", "全部数据"],
        ["售前", "自己提交的 + 作为关联售前的项目"],
        ["销售", "关联销售是自己的项目"],
        ["FAE 测试工程师", "分配给自己的项目"],
    ],
    widths=[8, 8]
)

add_heading("3.4 按钮级操作权限（项目详情）", 2)
add_table(
    ["操作", "可操作角色"],
    [
        ["开始测试 / 暂停 / 恢复 / 完成", "管理员 + 被分配的 FAE 测试人员"],
        ["更新中标", "所有人"],
        ["设为重点", "管理员 + 测试审批组"],
        ["填写进展 / 上传报告", "被分配的 FAE 测试人员 + 管理员"],
        ["删除项目", "管理员 + 资源管理员 + FAE 负责人"],
        ["预览项目", "所有可见该项目的用户（只读）"],
    ],
    widths=[7, 9]
)
add_para("所有权限控制均为前后端双重校验：前端按权限隐藏按钮，服务端接口再次校验，防止绕过。", indent=True)

# ===================== 4 核心业务流程 =====================
add_heading("四、核心业务流程", 1)
add_heading("4.1 测试申请审批流程", 2)
add_code(
"""售前工程师发起测试申请(填写申请表)
        │  提交
        ▼
  测试审批组审批 ──驳回(费用超标/需求不合理)──► 已驳回(可修改重提)
        │ 通过
        ▼
  是否借测超90天? ──是──► 分管领导二级审批 ──通过──┐
        │ 否                                       │
        ▼                                          ▼
  流转 FAE 负责人(daihw)分配任务 ◄───────────────────┘
        │  指定 FAE 测试工程师 + 资源类型
        ▼
  生成测试项目(未开始)
        │
        ▼
  未开始 → 进行中 ⇄ 暂停 → 已完成
        │  资源回收 → 报告编写/上传
        ▼
  任意节点可"关闭"(需填原因)"""
)
add_para("流程规则已固化进系统：", indent=True)
for t in [
    "申请时校验“暂停借用”三条件（超期 90 天借用 / 借用金额超销售任务 2% / 部门借用费用超部门任务 2%），命中则禁止提交；",
    "借测周期超 90 天自动进入分管领导二级审批，通过后流转 SPM 申请；",
    "循环机 / SPM 设备自动计入对应部门测试机费用台账，线上资源标记为公共资源不计费；",
    "借出超 30 天系统每两周自动生成跟催任务并提醒；超 12 个月未归还自动生成考核预警单。",
]:
    add_para("• " + t, indent=True)

# ===================== 5 功能模块 =====================
add_heading("五、功能模块清单", 1)
add_table(
    ["模块", "功能说明"],
    [
        ["项目统计", "多维统计看板（状态/类型/设备/周期/区域/月度趋势/人员排行/资源/中标），每个图表位置可自定义切换维度，支持数据权限隔离"],
        ["测试申请", "售前在线发起申请（必填校验），列表查询，草稿、撤回、驳回重提"],
        ["审批中心", "待办审批，点击项目名查看申请详情，通过/驳回/分配任务"],
        ["项目清单", "完整项目台账（覆盖原 24 字段），多维筛选、重点项目、删除、导出 Excel"],
        ["项目详情", "项目信息、状态流转、中标更新、进展时间线、测试报告管理，按钮级权限"],
        ["资源管理", "资源池台账、借用/归还闭环、超期预警（仅审批组/FAE负责人/资源管理员可见）"],
        ["消息通知", "站内信，可点击跳转到对应流程页面，未读角标"],
        ["系统管理", "用户管理、用户组管理、数据字典管理"],
    ],
    widths=[3, 13]
)

# ===================== 6 数据模型 =====================
add_heading("六、数据模型", 1)
add_para("核心业务表设计如下：", indent=True)
add_table(
    ["表名", "说明", "关键字段"],
    [
        ["sys_user / sys_role / sys_user_role", "用户/角色/用户角色", "username, oa_account(预留OA), email"],
        ["sys_user_group / sys_user_group_rel", "用户组/组成员", "group_code, leader_id(负责人)"],
        ["sys_dict / sys_config", "数据字典/系统参数", "dict_type, config_key(借用规则阈值)"],
        ["test_project", "测试项目主表", "客户/项目/区域/SPM/销售/售前/测试/类型/状态/中标/重点"],
        ["test_application", "测试申请单", "申请内容/current_node(流程节点)/status"],
        ["approval_record", "审批记录", "节点/审批人/动作/意见/耗时"],
        ["test_progress", "测试进展日报", "project_id, progress_date, content"],
        ["resource / resource_loan", "资源池/借用记录", "resource_type, 出厂价, 借用/归还时间, 费用归属"],
        ["test_report", "测试报告", "file_key(MinIO), version, upload_by"],
        ["notify_msg / audit_log", "站内通知/审计日志", "jump_url(跳转), is_read"],
    ],
    widths=[5.5, 4, 6.5]
)

# ===================== 7 技术选型 =====================
add_heading("七、技术选型", 1)
add_table(
    ["层次", "技术栈", "版本"],
    [
        ["前端", "Vue3 + Element Plus + Vite + ECharts + Pinia", "Vue 3.4 / Element Plus 2.7"],
        ["后端", "Java 17 + Spring Boot 3 + MyBatis-Plus + Spring Security", "Spring Boot 3.2.5"],
        ["数据库", "MySQL（utf8mb4）", "8.0"],
        ["缓存", "Redis", "7"],
        ["对象存储", "MinIO", "latest"],
        ["认证", "本地账号 + JWT（预留 OAuth2/CAS）", "jjwt 0.12.5"],
        ["部署", "Docker + Docker Compose + Nginx", "Docker 26"],
    ],
    widths=[3, 9, 4]
)

# ===================== 8 部署方案 =====================
add_heading("八、部署方案（已实现）", 1)

add_heading("8.1 部署架构", 2)
add_para("平台以 Docker Compose 单机编排方式部署于 192.168.101.165 虚拟机，共 5 个容器，仅 Nginx 暴露 80 端口对外：", indent=True)
add_code(
"""                    ┌────────────────────────────────────┐
   内网用户浏览器    │      192.168.101.165 (Rocky 8.10)    │
   ─────────────►  │  Docker Compose 编排                 │
      HTTP :80     │                                      │
                    │   tp-frontend (Nginx:80 对外)        │
                    │        │ /api 反代                    │
                    │   tp-backend  (Spring Boot:8080)     │
                    │      │      │        │               │
                    │   tp-mysql  tp-redis  tp-minio        │
                    │   (3306)   (6379)    (9000/9001)     │
                    │  ── 以上数据服务仅容器内网可达 ──       │
                    │  数据卷: /opt/testplatform/data/...   │
                    └────────────────────────────────────┘""")

add_heading("8.2 容器清单", 2)
add_table(
    ["容器", "镜像", "端口", "说明"],
    [
        ["tp-frontend", "testplatform-frontend（自研构建）", "80 → 对外", "前端静态资源 + Nginx 反向代理"],
        ["tp-backend", "testplatform-backend（自研构建）", "8080（内网）", "Spring Boot 后端服务"],
        ["tp-mysql", "mysql:8.0", "3306（内网）", "业务数据库"],
        ["tp-redis", "redis:7-alpine", "6379（内网）", "会话与缓存"],
        ["tp-minio", "minio/minio:latest", "9000/9001（内网）", "测试报告对象存储"],
    ],
    widths=[3, 5.5, 3, 5]
)

add_heading("8.3 服务器与目录", 2)
add_table(
    ["项目", "配置"],
    [
        ["操作系统", "Rocky Linux 8.10"],
        ["硬件", "20 核 CPU / 3.5GB 内存 / 96GB 磁盘（已用约 17GB）"],
        ["部署目录", "/opt/testplatform"],
        ["数据目录", "/opt/testplatform/data（mysql/redis/minio/logs/backup）"],
        ["编排文件", "/opt/testplatform/docker-compose.yml + .env"],
    ],
    widths=[4, 12]
)

add_heading("8.4 目录结构", 2)
add_code(
"""/opt/testplatform/
├── docker-compose.yml        # 一键编排
├── .env                      # 密码/密钥配置(权限600)
├── backend/                  # 后端源码 + Dockerfile
├── frontend/                 # 前端源码 + Dockerfile + nginx.conf
├── deploy/schema.sql         # 数据库初始化脚本
└── data/                     # 持久化数据
    ├── mysql/  redis/  minio/
    ├── backend/logs/
    └── backup/""")

add_heading("8.5 常用运维命令", 2)
add_code(
"""cd /opt/testplatform
docker compose ps                  # 查看容器状态
docker compose logs -f backend     # 查看后端日志
docker compose restart backend     # 重启后端
docker compose up -d               # 启动全部服务
docker compose build backend       # 重新构建后端镜像""")

add_heading("8.6 发版流程", 2)
for t in [
    "本地修改代码 → 上传服务器对应源码目录；",
    "执行 docker compose build backend（或 frontend）重新构建镜像；",
    "执行 docker compose up -d backend（或 frontend）滚动重启；",
    "数据库结构变更通过 schema.sql 增量脚本执行。",
]:
    add_para("• " + t, indent=True)

# ===================== 9 数据迁移 =====================
add_heading("九、历史数据迁移", 1)
add_para("已将原《应用测试部 2026 年测试项目汇总.xlsx》中的历史数据完整迁移至平台：", indent=True)
add_table(
    ["迁移内容", "数量", "说明"],
    [
        ["测试项目", "175 个", "完整覆盖原表 24 个统计字段，状态分布与原表一致"],
        ["测试进展", "569 条", "原进展列按日期自动拆分为日报记录"],
        ["中标统计", "7644 万元", "与原表口径一致"],
        ["重点项目", "可标记", "原重点测试项目表可手动关联标记"],
    ],
    widths=[4, 3, 9]
)
add_para("迁移完成后，飞书在线表格已冻结只读，新数据一律在平台产生。", indent=True)

# ===================== 10 安全与权限 =====================
add_heading("十、安全设计", 1)
for t in [
    "认证：本地账号 + BCrypt 密码加密 + JWT 无状态令牌（12 小时过期），首次登录强制改密；",
    "接口鉴权：JWT 过滤器解析用户，角色缓存 Redis 并可从数据库回源，缓存失效不影响鉴权；",
    "数据权限：按角色在申请/项目/统计各模块统一隔离；",
    "操作权限：按钮级权限前后端双重校验，防止接口绕过；",
    "模块权限：资源管理等敏感模块仅指定角色可见，菜单/路由/接口三层防护；",
    "审计：审批、驳回、分配、删除等关键操作全部留痕；",
    "网络安全：仅 80 端口对外，MySQL/Redis/MinIO 均不暴露宿主机。",
]:
    add_para("• " + t, indent=True)

# ===================== 11 后续规划 =====================
add_heading("十一、后续规划（二期/三期）", 1)
add_table(
    ["阶段", "内容"],
    [
        ["二期（资源与报表）", "资源池/借用规则引擎完善、超期预警、周报自动生成、企业微信/邮件通知联调"],
        ["三期（集成）", "OA 用户认证对接（OAuth2/CAS）、SPM 系统对接、BI 报表深化"],
    ],
    widths=[4, 12]
)
add_para("平台预留了 oa_account 字段与认证适配器接口，对接 OA 时仅需切换认证源，不影响业务数据。", indent=True)

# ===================== 附录 =====================
add_heading("附录：平台账号", 1)
add_table(
    ["账号", "姓名", "角色", "初始密码"],
    [
        ["admin", "系统管理员", "系统管理员", "Admin@123"],
        ["daihw", "戴海伟", "FAE负责人/资源管理员/测试审批组", "—"],
        ["test_sp", "审批组", "测试审批组", "—"],
        ["presale1", "售前测试1", "售前工程师", "—"],
        ["saletest1", "销售测试", "销售", "—"],
        ["faedemo", "FAE测试工程师", "FAE测试工程师", "—"],
    ],
    widths=[3.5, 4, 6, 2.5]
)
add_para("注：标“—”的账号密码由使用人自行保管，可在登录后通过“修改密码”自行修改。", indent=True, size=9)

doc.save(r"C:\Users\charl\Desktop\测试项目管理平台-开发设计部署方案.docx")
print("文档已生成: C:\\Users\\charl\\Desktop\\测试项目管理平台-开发设计部署方案.docx")
