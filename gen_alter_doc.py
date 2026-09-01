# -*- coding: utf-8 -*-
"""生成 alter.sql 导出与导入操作流程文档"""
import sys
sys.stdout.reconfigure(encoding="utf-8", errors="replace")
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

def set_cn(run, name="宋体", size=10.5, bold=False, color=None):
    run.font.name=name; run.font.size=Pt(size); run.font.bold=bold
    if color: run.font.color.rgb=color
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)

def H(text, level=1):
    p=doc.add_heading(level=level); r=p.add_run(text)
    sizes={1:18,2:14,3:12}
    r.font.name="微软雅黑"; r.font.size=Pt(sizes.get(level,12)); r.font.bold=True
    r.font.color.rgb=RGBColor(0xC0,0x1E,0x3A) if level==1 else RGBColor(0x1A,0x2A,0x6C)
    r._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')

def P(text, indent=False, size=10.5, bold=False):
    p=doc.add_paragraph()
    if indent: p.paragraph_format.first_line_indent=Cm(0.74)
    p.paragraph_format.space_after=Pt(4)
    r=p.add_run(text); set_cn(r,size=size,bold=bold)

def CODE(text):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(text); r.font.name="Consolas"; r.font.size=Pt(9)
    r._element.rPr.rFonts.set(qn('w:eastAsia'),'Consolas')
    pPr=p._p.get_or_add_pPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),'F2F2F2'); pPr.append(shd)

def TABLE(headers, rows, widths=None):
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=''
        r=c.paragraphs[0].add_run(h); set_cn(r,name="微软雅黑",size=10,bold=True,color=RGBColor(0xFF,0xFF,0xFF))
        shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),'1A2A6C'); c._tc.get_or_add_tcPr().append(shd)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=''
            r=cells[i].paragraphs[0].add_run(str(v)); set_cn(r,size=9.5)
    if widths:
        for i,w in enumerate(widths):
            for row in t.rows: row.cells[i].width=Cm(w)
    doc.add_paragraph()

# ===== 封面 =====
for _ in range(4): doc.add_paragraph()
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=t.add_run("数据库结构变更脚本(alter.sql)"); set_cn(r,name="微软雅黑",size=26,bold=True,color=RGBColor(0xC0,0x1E,0x3A))
s=doc.add_paragraph(); s.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=s.add_run("导出与正式平台导入 操作流程"); set_cn(r,name="微软雅黑",size=18,bold=True,color=RGBColor(0x1A,0x2A,0x6C))
doc.add_paragraph()
sub=doc.add_paragraph(); sub.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=sub.add_run("测试项目管理平台  |  2026-08-31 版本变更"); set_cn(r,size=12,color=RGBColor(0x88,0x88,0x88))
for _ in range(8): doc.add_paragraph()
info=doc.add_paragraph(); info.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=info.add_run("中科曙光  战略客户服务部\n\n2026 年 8 月"); set_cn(r,name="微软雅黑",size=14,color=RGBColor(0x33,0x33,0x33))
doc.add_page_break()

# ===== 1 概述 =====
H("一、alter.sql 是什么", 1)
P("alter.sql 是「数据库结构增量变更脚本」，记录某次版本升级中对数据库所做的结构性修改（新增字段、新增表、初始化字典/参数等 DDL/DML 语句）。", indent=True)
P("它的作用是：正式平台升级时，先在数据库上执行该脚本，使正式库的表结构与开发库保持一致，新版本代码才能正常运行。", indent=True)
P("alter.sql 由开发人员在功能开发时编写，存放在源码 deploy/ 目录，命名建议 alter_日期.sql（如 alter_20260831.sql），随版本一同发布。", indent=True)

H("二、本次(2026-08-31)的变更内容", 1)
P("本次开发新增「项目阶段(L1-L9)」和「招标状态」功能，涉及以下数据库变更：", indent=True)
TABLE(["序号","变更对象","变更内容"],
[
 ["1","test_application 表","新增 project_stage（项目阶段）、bid_status（招标状态）字段"],
 ["2","test_project 表","新增 project_stage（项目阶段）字段"],
 ["3","sys_dict 表","新增项目阶段字典 project_stage（L1~L9）"],
], widths=[1.5,5,9.5])

# ===== 2 导出 =====
H("三、alter.sql 的编写与导出（开发平台）", 1)

H("3.1 编写原则", 2)
for t in [
 "幂等可重复执行：所有 ALTER/INSERT 语句加存在性判断，重复执行不报错；",
 "新增字段：先查 information_schema.COLUMNS 判断字段是否已存在；",
 "字典/参数初始化：INSERT 用 ON DUPLICATE KEY UPDATE；",
 "中文内容统一使用 utf8mb4；",
 "文件末尾给出执行完成提示。",
]:
 P("• "+t, indent=True)

H("3.2 本次 alter.sql 完整内容", 2)
P("本次变更脚本已编写完成并存放于开发平台 /opt/testplatform/deploy/alter_20260831.sql，内容如下：", indent=True)
CODE("""-- 测试项目管理平台 数据库增量变更脚本  v2026-08-31
-- 变更内容: 新增项目阶段(L1-L9)与招标状态
-- 说明: 本脚本幂等,可重复执行不报错

-- 1. 申请表新增 项目阶段、招标状态
SET @app_has_stage := (SELECT COUNT(*) FROM information_schema.COLUMNS
  WHERE TABLE_SCHEMA='test_platform' AND TABLE_NAME='test_application' AND COLUMN_NAME='project_stage');
SET @sql := IF(@app_has_stage=0,
  'ALTER TABLE test_application ADD COLUMN project_stage VARCHAR(20) NULL COMMENT ''项目阶段L1-L9'' AFTER spm_no',
  'SELECT ''test_application.project_stage 已存在,跳过'' AS tip');
PREPARE stmt FROM @sql; EXECUTE stmt; DEALLOCATE PREPARE stmt;

SET @app_has_bid := (SELECT COUNT(*) FROM information_schema.COLUMNS
  WHERE TABLE_SCHEMA='test_platform' AND TABLE_NAME='test_application' AND COLUMN_NAME='bid_status');
SET @sql := IF(@app_has_bid=0,
  'ALTER TABLE test_application ADD COLUMN bid_status VARCHAR(20) NULL COMMENT ''招标状态'' AFTER project_stage',
  'SELECT ''test_application.bid_status 已存在,跳过'' AS tip');
PREPARE stmt FROM @sql; EXECUTE stmt; DEALLOCATE PREPARE stmt;

-- 2. 项目表新增 项目阶段
SET @proj_has_stage := (SELECT COUNT(*) FROM information_schema.COLUMNS
  WHERE TABLE_SCHEMA='test_platform' AND TABLE_NAME='test_project' AND COLUMN_NAME='project_stage');
SET @sql := IF(@proj_has_stage=0,
  'ALTER TABLE test_project ADD COLUMN project_stage VARCHAR(20) NULL COMMENT ''项目阶段L1-L9'' AFTER spm_no',
  'SELECT ''test_project.project_stage 已存在,跳过'' AS tip');
PREPARE stmt FROM @sql; EXECUTE stmt; DEALLOCATE PREPARE stmt;

-- 3. 新增项目阶段字典(L1-L9)
INSERT INTO sys_dict (dict_type, dict_label, dict_value, sort) VALUES
('project_stage','L1','L1',1),('project_stage','L2','L2',2),('project_stage','L3','L3',3),
('project_stage','L4','L4',4),('project_stage','L5','L5',5),('project_stage','L6','L6',6),
('project_stage','L7','L7',7),('project_stage','L8','L8',8),('project_stage','L9','L9',9)
ON DUPLICATE KEY UPDATE dict_label=VALUES(dict_label), sort=VALUES(sort);

SELECT 'alter.sql 执行完成' AS result;""")

H("3.3 导出/获取脚本文件", 2)
P("脚本已存放于开发平台源码目录，可直接下载：", indent=True)
CODE("""# 开发平台上脚本位置
/opt/testplatform/deploy/alter_20260831.sql

# 从开发平台下载到本地(本地执行)
scp root@192.168.101.165:/opt/testplatform/deploy/alter_20260831.sql ./""")

# ===== 3 导入 =====
H("四、正式平台导入流程", 1)

H("4.1 导入前准备（必须先备份）", 2)
P("执行结构变更前，务必先完成数据备份：", indent=True)
CODE("""# 在正式服务器执行备份
bash /opt/testplatform/backup.sh""")

H("4.2 上传脚本到正式平台", 2)
CODE("""# 本地执行,上传脚本到正式服务器 deploy 目录
scp alter_20260831.sql root@<正式服务器IP>:/opt/testplatform/deploy/""")

H("4.3 执行导入", 2)
CODE("""# 在正式服务器执行
cd /opt/testplatform
# 读取数据库密码(并去除配置文件回车符)
set -a; source .env; set +a
DB_PASSWORD=$(echo -n "$DB_PASSWORD" | tr -d '\\r')

# 执行结构变更脚本
docker exec -i tp-mysql mysql -utestplatform -p"$DB_PASSWORD" \\
  --default-character-set=utf8mb4 test_platform < deploy/alter_20260831.sql""")

H("4.4 验证变更结果", 2)
CODE("""# 验证新增字段是否生效
docker exec -i tp-mysql mysql -utestplatform -p"$DB_PASSWORD" --default-character-set=utf8mb4 test_platform -e "
SHOW COLUMNS FROM test_application LIKE 'project_stage';
SHOW COLUMNS FROM test_application LIKE 'bid_status';
SHOW COLUMNS FROM test_project LIKE 'project_stage';
SELECT dict_label, dict_value FROM sys_dict WHERE dict_type='project_stage' ORDER BY sort;
\"""")
P("预期结果：三个字段均存在，字典返回 L1~L9 共 9 条记录。", indent=True)

H("4.5 导入后继续升级", 2)
P("数据库结构变更验证通过后，再进行应用代码升级（重新构建并重启前后端），完成本次版本升级。", indent=True)

# ===== 5 完整流程图 =====
H("五、完整流程一览", 1)
CODE("""【开发平台】
  功能开发 -> 编写 alter_日期.sql(幂等) -> 存入 deploy/ 目录
      │
  scp 导出脚本
      ▼
【正式平台】
  1. 数据备份  bash /opt/testplatform/backup.sh
  2. 上传脚本  scp alter_xxx.sql -> deploy/
  3. 执行导入  docker exec ... mysql < deploy/alter_xxx.sql
  4. 验证变更  SHOW COLUMNS / SELECT 字典
  5. 升级应用  docker compose build + up -d
  6. 验证功能""")

H("六、注意事项", 1)
for t in [
 "顺序：先备份 → 再执行 alter.sql → 再升级应用代码，顺序不可颠倒；",
 "alter.sql 必须幂等，允许重复执行；本次脚本已做幂等处理并实测通过；",
 "正式平台数据库密码从 .env 读取，注意去除 CRLF 回车符；",
 "若执行报错，先核对报错信息，必要时用备份回滚后再排查；",
 "每次结构变更都应记录到 deploy/ 目录并纳入版本管理。",
]:
 P("• "+t, indent=True)

doc.save(r"C:\Users\charl\Desktop\alter.sql导出与导入操作流程.docx")
print("已生成: C:\\Users\\charl\\Desktop\\alter.sql导出与导入操作流程.docx")
