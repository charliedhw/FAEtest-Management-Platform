# -*- coding: utf-8 -*-
"""生成 测试项目管理平台 迁移与升级方案 文档"""
import sys
sys.stdout.reconfigure(encoding="utf-8", errors="replace")
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

def set_cn(run, name="宋体", size=10.5, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color: run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)

def H(text, level=1):
    p = doc.add_heading(level=level)
    r = p.add_run(text)
    sizes = {1:18,2:14,3:12}
    r.font.name="微软雅黑"; r.font.size=Pt(sizes.get(level,12)); r.font.bold=True
    r.font.color.rgb = RGBColor(0xC0,0x1E,0x3A) if level==1 else RGBColor(0x1A,0x2A,0x6C)
    r._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')

def P(text, indent=False, size=10.5, bold=False):
    p = doc.add_paragraph()
    if indent: p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); set_cn(r, size=size, bold=bold)

def CODE(text):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.font.name="Consolas"; r.font.size=Pt(9)
    r._element.rPr.rFonts.set(qn('w:eastAsia'),'Consolas')
    pPr = p._p.get_or_add_pPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'),'F2F2F2'); pPr.append(shd)

def TABLE(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=''
        r=c.paragraphs[0].add_run(h); set_cn(r,name="微软雅黑",size=10,bold=True,color=RGBColor(0xFF,0xFF,0xFF))
        shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),'1A2A6C'); c._tc.get_or_add_tcPr().append(shd)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=''
            r=cells[i].paragraphs[0].add_run(str(v)); set_cn(r,size=9.5)
    if widths:
        for i,w in enumerate(widths):
            for row in t.rows: row.cells[i].width=Cm(w)
    doc.add_paragraph()

# ===== 封面 =====
for _ in range(4): doc.add_paragraph()
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=t.add_run("测试项目管理平台"); set_cn(r,name="微软雅黑",size=30,bold=True,color=RGBColor(0xC0,0x1E,0x3A))
s=doc.add_paragraph(); s.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=s.add_run("服务器迁移与系统升级方案"); set_cn(r,name="微软雅黑",size=20,bold=True,color=RGBColor(0x1A,0x2A,0x6C))
doc.add_paragraph()
sub=doc.add_paragraph(); sub.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=sub.add_run("—— 含数据备份、迁移实施、安全升级全流程 ——"); set_cn(r,size=12,color=RGBColor(0x88,0x88,0x88))
for _ in range(8): doc.add_paragraph()
info=doc.add_paragraph(); info.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=info.add_run("中科曙光  战略客户服务部\n\n2026 年 8 月"); set_cn(r,name="微软雅黑",size=14,color=RGBColor(0x33,0x33,0x33))
doc.add_page_break()

# ===== 目录说明 =====
H("文档说明", 1)
P("本文档包含两部分：", indent=True)
P("第一部分：平台从当前开发/测试服务器（192.168.101.165）迁移到正式 Linux 服务器的完整迁移方案；", indent=True)
P("第二部分：平台功能开发完毕后，在生产环境进行版本升级的标准升级方案（升级前必须先完成数据备份）。", indent=True)
doc.add_page_break()

# ============================================================
# 第一部分 迁移方案
# ============================================================
H("第一部分  平台迁移方案（迁移到正式 Linux 服务器）", 1)

H("1.1 迁移概述", 2)
P("当前平台部署于 192.168.101.165（Rocky Linux 8.10）。迁移目标为一台正式 Linux 服务器。平台采用全容器化（Docker）部署，迁移本质是“镜像 + 数据卷 + 配置”的整体搬迁，具有良好可移植性。", indent=True)

H("1.2 迁移范围", 2)
TABLE(["迁移项","内容","位置/方式"],
[
 ["应用镜像","frontend / backend 自研镜像","docker save 导出为 tar 包"],
 ["数据库","MySQL 业务数据（test_platform 库）","mysqldump 导出 / 数据卷拷贝"],
 ["文件存储","MinIO 中的测试报告文件","数据卷打包"],
 ["缓存","Redis（会话/缓存，非关键）","可不迁移，重启后自动重建"],
 ["配置文件","docker-compose.yml / .env / nginx","直接拷贝"],
 ["初始化脚本","deploy/schema.sql","随配置拷贝"],
], widths=[3,7,6])

H("1.3 目标服务器准备", 2)
TABLE(["项目","最低要求","说明"],
[
 ["操作系统","Rocky Linux 8 / Ubuntu 22.04 / 麒麟V10","与源环境一致最佳"],
 ["CPU / 内存","4 核 / 8 GB","推荐 8 核 / 16 GB"],
 ["磁盘","200 GB SSD","含数据增长与备份空间"],
 ["Docker","20.10+ 及 docker compose 插件","参考安装脚本"],
 ["网络","内网可达，开放 80 端口","如有域名更佳"],
], widths=[3.5,5.5,7])

P("目标服务器安装 Docker（若未安装）：", indent=True)
CODE("""# 安装 Docker（以阿里云镜像为例）
dnf install -y yum-utils
yum-config-manager --add-repo https://mirrors.aliyun.com/docker-ce/linux/centos/docker-ce.repo
sed -i 's|download.docker.com|mirrors.aliyun.com/docker-ce|g' /etc/yum.repos.d/docker-ce.repo
dnf install -y --allowerasing docker-ce docker-ce-cli containerd.io docker-compose-plugin
systemctl enable --now docker

# 配置镜像加速与 DNS
mkdir -p /etc/docker
cat > /etc/docker/daemon.json <<'EOF'
{
  "registry-mirrors": ["https://docker.m.daocloud.io", "https://dockerproxy.net"],
  "dns": ["114.114.114.114", "8.8.8.8"],
  "log-driver": "json-file",
  "log-opts": {"max-size": "50m", "max-file": "3"}
}
EOF
systemctl restart docker""")

H("1.4 迁移实施步骤", 2)

P("步骤一：源服务器数据备份", 3)
P("直接运行平台内置备份脚本（与升级备份共用同一脚本，详见第二部分 2.2 节）：", indent=True)
CODE("""# 在源服务器 192.168.101.165 执行
bash /opt/testplatform/backup.sh
# 备份产物位于 /opt/testplatform/data/backup/
#  - db_test_platform_时间戳.sql.gz   (数据库)
#  - minio_时间戳.tar.gz              (报告文件)
#  - config_时间戳.tar.gz             (配置)
# 打包备份产物便于传输
tar -czf /root/tp_backup_$(date +%Y%m%d).tar.gz -C /opt/testplatform/data/backup .""")

P("步骤二：导出应用镜像", 3)
CODE("""# 在源服务器导出前后端镜像
docker save testplatform-backend testplatform-frontend | gzip > /root/tp_images.tar.gz""")

P("步骤三：传输到目标服务器", 3)
CODE("""# 将备份包与镜像包拷贝到目标服务器
scp /root/tp_backup_*.tar.gz root@<目标IP>:/root/
scp /root/tp_images.tar.gz root@<目标IP>:/root/""")

P("步骤四：目标服务器部署", 3)
CODE("""# 在目标服务器执行
mkdir -p /opt/testplatform/data/backup && cd /opt/testplatform
# 1. 解包备份产物(数据库/报告/配置)到备份目录
tar -xzf /root/tp_backup_*.tar.gz -C /opt/testplatform/data/backup/
# 2. 从备份中取出配置文件(.env / docker-compose.yml / deploy)到部署目录
tar -xzf /opt/testplatform/data/backup/config_*.tar.gz -C /opt/testplatform/
# 3. 导入镜像
docker load < /root/tp_images.tar.gz
# 4. 恢复 MinIO 报告文件到数据目录
tar -xzf /opt/testplatform/data/backup/minio_*.tar.gz -C /opt/testplatform/data/
# 5. 拷贝源码目录(用于后续构建升级,可选但推荐)
#    将源服务器 backend/ frontend/ 目录整体拷贝到 /opt/testplatform/

# 6. 修改 .env 中与环境相关的配置(如数据库密码、域名等)
# 7. 先启动基础服务
docker compose up -d mysql redis minio""")

P("步骤五：恢复数据库", 3)
CODE("""# 备份的数据库是 .sql.gz 压缩格式,需先解压再导入
cd /opt/testplatform/data/backup
gunzip -c db_test_platform_*.sql.gz > /tmp/restore.sql

# 等待 mysql 健康后导入(使用应用账号 testplatform 或 root)
docker exec -i tp-mysql mysql -utestplatform -p<DB_PASSWORD> \\
  --default-character-set=utf8mb4 test_platform < /tmp/restore.sql

# 验证数据
docker exec -i tp-mysql mysql -utestplatform -p<DB_PASSWORD> test_platform \\
  -e "SELECT COUNT(*) FROM test_project;""")

P("步骤六：启动应用并验证", 3)
CODE("""docker compose up -d
docker compose ps
# 浏览器访问 http://<目标IP> 验证登录与各模块功能""")

H("1.5 迁移验证清单", 2)
TABLE(["验证项","预期结果"],
[
 ["容器状态","5 个容器全部 Up，mysql healthy"],
 ["登录","admin 可正常登录"],
 ["项目清单","175 个项目完整显示"],
 ["统计数据","看板各维度数据正确"],
 ["测试报告","可正常上传/下载"],
 ["消息通知","通知可点击跳转"],
 ["中文显示","无乱码"],
], widths=[5,11])

H("1.6 回退方案", 2)
P("若迁移后验证不通过，直接切换回源服务器继续提供服务即可（源服务器数据未删除）。目标服务器问题排查修复后重新迁移。建议在非工作时间执行迁移，预留回退窗口。", indent=True)

doc.add_page_break()

# ============================================================
# 第二部分 升级方案
# ============================================================
H("第二部分  系统升级方案（升级前必须数据备份）", 1)

H("2.1 升级原则", 2)
for t in [
 "升级前必须先完成数据备份，备份未成功不得升级；",
 "先备份数据库，再备份文件，最后才是应用升级；",
 "升级操作建议在业务低峰期（如晚间）执行；",
 "每次升级保留回退能力（旧镜像 + 完整备份）。",
]:
 P("• "+t, indent=True)

H("2.2 升级前数据备份（必须步骤）", 2)
P("平台已内置标准化备份脚本 /opt/testplatform/backup.sh（已实测验证），升级前执行一次即可完成数据库与文件的全量备份。", indent=True)

P("备份脚本功能与备份内容：", 3)
TABLE(["备份项","文件格式","说明"],
[
 ["MySQL 数据库","db_test_platform_时间戳.sql.gz","test_platform 全库（表结构+数据+索引+触发器）"],
 ["MinIO 报告文件","minio_时间戳.tar.gz","所有测试报告附件"],
 ["配置文件","config_时间戳.tar.gz",".env / docker-compose.yml / schema.sql"],
 ["备份日志","backup_时间戳.log","每次备份的执行日志"],
], widths=[4,5.5,6.5])

P("备份脚本使用方法：", 3)
CODE("""# 手动备份(默认保留30天)
bash /opt/testplatform/backup.sh

# 手动备份并指定保留天数(如保留60天)
bash /opt/testplatform/backup.sh 60""")

P("备份脚本特性：", 3)
for t in [
 "自动读取 .env 中的数据库密码，并兼容处理配置文件中的回车符；",
 "备份有效性自动校验：数据库备份后检查是否包含 CREATE TABLE，空文件会报错退出；",
 "自动清理过期备份文件，默认保留 30 天（可通过参数调整）；",
 "每次备份生成带时间戳的执行日志，便于追溯。",
]:
 P("• "+t, indent=True)

P("定时自动备份（已配置）：", 3)
P("系统已通过 crontab 配置每天凌晨 2:00 自动备份，保留 30 天，无需人工干预：", indent=True)
CODE("""# 已配置的定时任务(crontab -l 可查看)
0 2 * * * /bin/bash /opt/testplatform/backup.sh 30 >> /opt/testplatform/data/backup/cron.log 2>&1""")

P("备份文件统一存放于 /opt/testplatform/data/backup/ 目录。备份结果验证：", indent=True)
CODE("""# 查看备份文件
ls -lh /opt/testplatform/data/backup/

# 验证数据库备份完整性(应能看到 CREATE TABLE 等正常 SQL 内容)
zcat /opt/testplatform/data/backup/db_test_platform_<时间戳>.sql.gz | head -20""")

H("2.3 升级操作步骤", 2)

P("步骤一：上传新版本代码", 3)
P("将开发完成的前后端源码上传到服务器对应目录（/opt/testplatform/backend、/opt/testplatform/frontend）。", indent=True)

P("步骤二：数据库结构变更（如有）", 3)

P("【alter.sql 是什么、在哪里】", 3)
P("alter.sql 是本次升级涉及的「数据库结构增量变更脚本」（新增字段、新增表、修改索引、初始化字典/参数等 DDL/DML 语句）。它不是系统自动生成的，而是由开发人员在本次功能开发过程中编写并随版本一起提供的，通常存放于源码的 deploy/ 目录下（如 deploy/alter_v1.x.sql），与部署脚本一同管理。", indent=True)

P("【如何编写/导出 alter.sql（开发平台）】", 3)
P("方式一（推荐）：开发过程中手工编写增量脚本。开发人员把本次改动的 SQL 语句写入 deploy/alter_版本号.sql，例如：", indent=True)
CODE("""-- deploy/alter_20260831.sql 示例
ALTER TABLE test_application ADD COLUMN project_stage VARCHAR(20) NULL COMMENT '项目阶段';
ALTER TABLE test_application ADD COLUMN bid_status VARCHAR(20) NULL COMMENT '招标状态';
INSERT INTO sys_dict (dict_type, dict_label, dict_value, sort) VALUES
('project_stage','L1','L1',1) ON DUPLICATE KEY UPDATE dict_label=VALUES(dict_label);""")
P("方式二：用 mysqldump 从开发库导出指定表结构（仅结构、不含数据）：", indent=True)
CODE("""# 在开发平台导出某张表的结构变更(示例:导出 test_application 表结构)
docker exec tp-mysql mysqldump -utestplatform -p<密码> \\
  --default-character-set=utf8mb4 --no-data test_platform test_application \\
  > alter_table.sql
# 注意: 此方式导出的是整表结构, 适用于新增表; 新增字段建议仍用方式一手工编写 ALTER 语句""")

P("【如何在正式平台导入 alter.sql】", 3)
P("先将 alter.sql 上传到正式服务器，再导入数据库：", indent=True)
CODE("""# 1. 上传脚本到正式服务器(本地执行)
scp deploy/alter_20260831.sql root@<正式服务器IP>:/opt/testplatform/deploy/

# 2. 在正式服务器导入(数据库密码从 .env 读取,注意去除回车符)
cd /opt/testplatform
set -a; source .env; set +a
DB_PASSWORD=$(echo -n "$DB_PASSWORD" | tr -d '\\r')
docker exec -i tp-mysql mysql -utestplatform -p"$DB_PASSWORD" \\
  --default-character-set=utf8mb4 test_platform < deploy/alter_20260831.sql

# 3. 验证变更是否生效(示例:查看新增字段)
docker exec -i tp-mysql mysql -utestplatform -p"$DB_PASSWORD" test_platform \\
  -e "SHOW COLUMNS FROM test_application LIKE 'project_stage';""")

P("注意事项：", 3)
for t in [
 "alter.sql 必须先于应用升级执行（先改库结构，再部署新代码），顺序不能反；",
 "执行前必须先完成数据备份（见 2.2 节），防止 DDL 执行异常导致数据问题；",
 "脚本中的语句建议全部加幂等保护（如 IF NOT EXISTS / ON DUPLICATE KEY UPDATE），重复执行不报错；",
 "执行后务必验证表结构变更是否生效，再进行应用升级。",
]:
 P("• "+t, indent=True)

P("步骤三：重新构建并升级应用", 3)
CODE("""cd /opt/testplatform
# 重新构建镜像(改动哪个就构建哪个)
docker compose build backend
docker compose build frontend
# 滚动重启(先后端后前端)
docker compose up -d backend
docker compose up -d frontend
# 查看状态与日志
docker compose ps
docker compose logs -f backend""")

P("步骤四：升级后验证", 3)
TABLE(["验证项","说明"],
[
 ["服务启动","docker compose ps 全部 Up，日志无 ERROR"],
 ["登录","各角色账号可正常登录"],
 ["核心流程","申请→审批→分配→进展→报告 走通"],
 ["新功能","本次升级的新功能点逐一验证"],
 ["数据","升级前数据完整无丢失"],
], widths=[4,12])

H("2.4 升级回退方案", 2)
P("若升级后验证不通过，按以下步骤回退：", indent=True)
CODE("""# 1. 回退代码与镜像(用旧镜像重新启动)
cd /opt/testplatform
docker compose up -d --force-recreate backend frontend

# 2. 如数据库结构变更导致问题,恢复备份
zcat data/backup/db_test_platform_<时间戳>.sql.gz | \\
  docker exec -i tp-mysql mysql -uroot -p${MYSQL_ROOT_PASSWORD} \\
  --default-character-set=utf8mb4 test_platform

# 3. 如文件存储异常,恢复 minio
tar -xzf data/backup/minio_<时间戳>.tar.gz -C /opt/testplatform/data/""")

H("2.5 升级 checklist（升级执行单）", 2)
TABLE(["序号","检查项","是否完成"],
[
 ["1","数据备份已完成且验证有效","□"],
 ["2","新版本代码已上传服务器","□"],
 ["3","数据库变更脚本已执行(如有)","□"],
 ["4","应用镜像已重新构建","□"],
 ["5","服务重启后全部 Up","□"],
 ["6","登录及核心流程验证通过","□"],
 ["7","新功能验证通过","□"],
 ["8","旧镜像/备份已保留可回退","□"],
], widths=[2,11,3])

H("2.6 注意事项", 2)
for t in [
 "数据库密码、JWT 密钥等敏感信息统一在 .env 管理，.env 不进版本库；",
 "升级涉及数据库变更时，务必先备份再执行 DDL；",
 "保留最近 30 天备份，超期备份自动清理（可配置定时任务）；",
 "每次升级记录升级时间、版本号、变更内容、操作人，便于追溯。",
]:
 P("• "+t, indent=True)

doc.save(r"C:\Users\charl\Desktop\测试项目管理平台-迁移与升级方案.docx")
print("已生成: C:\\Users\\charl\\Desktop\\测试项目管理平台-迁移与升级方案.docx")
